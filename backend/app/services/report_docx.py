"""Render the report Markdown subset to a customer-ready Word document."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor, Twips

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_HEADING_PATTERN = re.compile(r"^(#{1,3})\s+(.*)$")
_LIST_PATTERN = re.compile(r"^[-*]\s+(.*)$")
_TABLE_SEPARATOR_CELL = re.compile(r"^:?-{3,}:?$")
_INLINE_BOLD_PATTERN = re.compile(r"(\*\*[^*]+\*\*)")

# standard_business_brief 的中文客户报告本地化：A4 是国内商务文档的通用纸型，
# 宋体/微软雅黑避免中文回落字体漂移。正文区宽 16.8 cm = 9524 DXA。
_TABLE_WIDTH_DXA = 9524
_TABLE_INDENT_DXA = 120
_TABLE_CELL_MARGINS_DXA = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def render_report_docx(markdown: str, *, title: str) -> bytes:
    document = Document()
    _configure_document(document, title)
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        table = _read_markdown_table(lines, index)
        if table is not None:
            rows, index = table
            _add_table(document, rows)
            continue

        heading_match = _HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text_value = heading_match.group(2).strip()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline_runs(paragraph, text_value)
            index += 1
            continue

        list_match = _LIST_PATTERN.match(stripped)
        if list_match:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, list_match.group(1).strip())
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            _set_paragraph_shading(paragraph, "F3F4F6")
            _add_inline_runs(paragraph, stripped.lstrip(">").strip(), color="4B5563")
            index += 1
            continue

        if stripped in {"---", "***"}:
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_value = lines[index].strip()
            if not next_value:
                break
            if (
                _HEADING_PATTERN.match(next_value)
                or _LIST_PATTERN.match(next_value)
                or next_value.startswith(">")
                or _read_markdown_table(lines, index) is not None
            ):
                break
            paragraph_lines.append(next_value)
            index += 1
        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, " ".join(paragraph_lines))

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def safe_docx_filename(title: str | None) -> str:
    value = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", str(title or "推荐报告")).strip(" ._")
    return f"{(value or '推荐报告')[:80]}.docx"


def _configure_document(document: Any, title: str) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    _set_east_asia_font(normal, "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "宋体"
    list_bullet.font.size = Pt(10.5)
    _set_east_asia_font(list_bullet, "宋体")
    list_bullet.paragraph_format.left_indent = Cm(1.27)
    list_bullet.paragraph_format.first_line_indent = Cm(-0.635)
    list_bullet.paragraph_format.space_after = Pt(8)
    list_bullet.paragraph_format.line_spacing = 1.167

    heading_settings = {
        "Heading 1": ("微软雅黑", 20, "1F2937", 14, 10),
        "Heading 2": ("微软雅黑", 15, "1F2937", 14, 6),
        "Heading 3": ("微软雅黑", 12, "374151", 10, 4),
    }
    for name, (font_name, size, color, before, after) in heading_settings.items():
        style = styles[name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        _set_east_asia_font(style, font_name)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    document.core_properties.title = title
    document.core_properties.subject = "Match-MA 推荐报告"
    document.core_properties.author = "Match-MA"


def _read_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or "|" not in lines[start] or "|" not in lines[start + 1]:
        return None
    header = _split_table_row(lines[start])
    separator = _split_table_row(lines[start + 1])
    if not header or len(header) != len(separator):
        return None
    if not all(_TABLE_SEPARATOR_CELL.match(cell.replace(" ", "")) for cell in separator):
        return None
    rows = [header]
    index = start + 2
    while index < len(lines) and "|" in lines[index]:
        row = _split_table_row(lines[index])
        if not row:
            break
        if len(row) < len(header):
            row.extend([""] * (len(header) - len(row)))
        rows.append(row[: len(header)])
        index += 1
    return rows, index


def _split_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|") and not value.endswith("\\|"):
        value = value[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for character in value:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            escaped = True
        elif character == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
    cells.append("".join(current).strip())
    return cells


def _add_table(document: Any, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = False
    widths = _column_widths(column_count)
    _set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        _prevent_row_split(row)
        if row_index == 0:
            _repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.width = Twips(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, "E5E7EB")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            _add_inline_runs(paragraph, value, bold=row_index == 0)
            for run in paragraph.runs:
                run.font.size = Pt(9)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_inline_runs(
    paragraph: Any,
    value: str,
    *,
    bold: bool = False,
    color: str | None = None,
) -> None:
    parts = _INLINE_BOLD_PATTERN.split(value)
    for part in parts:
        if not part:
            continue
        inline_bold = part.startswith("**") and part.endswith("**")
        text_value = part[2:-2] if inline_bold else part
        run = paragraph.add_run(text_value)
        run.bold = bold or inline_bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _set_east_asia_font(style: Any, font_name: str) -> None:
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_cell_shading(cell: Any, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_shading(paragraph: Any, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shading)


def _repeat_table_header(row: Any) -> None:
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(table_header)


def _prevent_row_split(row: Any) -> None:
    cant_split = OxmlElement("w:cantSplit")
    row._tr.get_or_add_trPr().append(cant_split)


def _column_widths(column_count: int) -> list[int]:
    if column_count == 4:
        return [1700, 2200, 3500, 2124]
    if column_count == 3:
        return [2200, 4000, 3324]
    if column_count == 2:
        return [3000, 6524]
    if column_count == 1:
        return [_TABLE_WIDTH_DXA]
    base = _TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += _TABLE_WIDTH_DXA - sum(widths)
    return widths


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.insert(0, table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(_TABLE_WIDTH_DXA))

    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(_TABLE_INDENT_DXA))

    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))


def _set_cell_margins(cell: Any) -> None:
    cell_pr = cell._tc.get_or_add_tcPr()
    margins = cell_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_pr.append(margins)
    for side, width in _TABLE_CELL_MARGINS_DXA.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")
