from __future__ import annotations

from io import BytesIO
from uuid import UUID

import pytest
from docx import Document

from backend.app.services import recommendation_report as report_service
from backend.app.services.recommendation_report import (
    BUYER_FACING_TARGET_REPORT,
    REPORT_MAX_ITEMS,
    SELLER_FACING_BUYER_REPORT,
    build_fallback_report_markdown,
    build_recommendation_report_context,
    default_report_type,
    ensure_report_item_count,
    normalize_report_markdown,
    report_type_matches_mode,
)
from backend.app.services.report_docx import render_report_docx, safe_docx_filename

SESSION_ID = UUID("10000000-0000-0000-0000-000000000001")
TARGET_ID = UUID("20000000-0000-0000-0000-000000000001")
INTENT_ID = UUID("30000000-0000-0000-0000-000000000001")
PARTY_ID = UUID("40000000-0000-0000-0000-000000000001")
ITEM_ID = UUID("50000000-0000-0000-0000-000000000001")


def test_report_types_are_direction_specific() -> None:
    assert default_report_type("buyer_to_target") == BUYER_FACING_TARGET_REPORT
    assert default_report_type("target_to_buyer") == SELLER_FACING_BUYER_REPORT
    assert report_type_matches_mode(BUYER_FACING_TARGET_REPORT, "buyer_to_target")
    assert not report_type_matches_mode(SELLER_FACING_BUYER_REPORT, "buyer_to_target")


def test_report_item_hard_limit_is_ten() -> None:
    ensure_report_item_count([{"id": index} for index in range(REPORT_MAX_ITEMS)])
    with pytest.raises(ValueError, match="at most 10"):
        ensure_report_item_count([{"id": index} for index in range(REPORT_MAX_ITEMS + 1)])
    with pytest.raises(ValueError, match="At least one"):
        ensure_report_item_count([])


def test_context_keeps_long_profile_sections_deep_eval_and_real_buyer_name(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    target_row = {
        "id": TARGET_ID,
        "target_name": "示例新能源标的",
        "business_summary": "结构化业务摘要",
        "current_revenue_yuan": 123_000_000,
    }
    intent_row = {
        "id": INTENT_ID,
        "intent_name": "新能源并购需求",
        "raw_requirement_text": "希望收购具备核心技术和稳定团队的新能源项目。",
        "intent_summary": "新能源控股型收购",
    }
    party_row = {"id": PARTY_ID, "buyer_name": "真实买家集团", "notes": "产业协同能力较强"}

    monkeypatch.setattr(
        report_service,
        "_load_entity_rows",
        lambda _db, entity_type, _ids: {
            str(TARGET_ID): target_row,
        } if entity_type == "seller_target" else {str(INTENT_ID): intent_row},
    )
    monkeypatch.setattr(
        report_service,
        "_load_buyer_parties",
        lambda _db, _ids: {str(PARTY_ID): party_row},
    )
    monkeypatch.setattr(
        report_service,
        "load_profile_sections",
        lambda _db, *, entity_type, entity_ids: {
            str(entity_ids[0]): {
                "business_product": {
                    "content_text": "深度画像中的大段业务与产品描述",
                    "info_status": "filled",
                }
            }
        },
    )
    monkeypatch.setattr(
        report_service,
        "_load_buyer_intent_scenarios",
        lambda _db, _ids: {str(INTENT_ID): [{"label": "控股方案", "fields_json": {}}]},
    )
    monkeypatch.setattr(
        report_service,
        "_load_latest_candidate_snapshots",
        lambda _db, _session_id: {
            f"{TARGET_ID}:{INTENT_ID}": {
                "evidence_json": {"matches": ["行业方向匹配"], "gaps": ["需确认客户集中度"]},
                "deep_eval": {
                    "reason": "现阶段具备继续接触价值",
                    "risks": ["客户集中度尚不明确"],
                    "info_gaps": ["近两年客户结构"],
                },
            }
        },
    )

    selected = [{
        "id": ITEM_ID,
        "seller_target_id": TARGET_ID,
        "buyer_intent_id": INTENT_ID,
        "buyer_party_id": PARTY_ID,
        "buyer_intent_name": "新能源并购需求",
        "seller_target_name": "示例新能源标的",
    }]
    session = {
        "id": SESSION_ID,
        "mode": "target_to_buyer",
        "initial_condition_snapshot_json": {},
        "latest_condition_snapshot_json": {},
        "condition_overrides_json": {},
    }
    context = build_recommendation_report_context(
        object(),
        report={"id": "", "report_type": SELLER_FACING_BUYER_REPORT, "title": "推荐买家报告"},
        session=session,
        selected_items=selected,
    )

    candidate = context["candidates"][0]
    assert candidate["candidate"]["buyer_party"]["buyer_name"] == "真实买家集团"
    assert "深度画像中的大段业务与产品描述" in str(candidate["candidate"])
    assert "希望收购具备核心技术" in str(candidate["candidate"])
    assert candidate["matching_context"]["deep_evaluation"]["risks"] == ["客户集中度尚不明确"]


def test_fallback_and_normalizer_stay_in_preview_contract() -> None:
    context = {
        "report": {"audience": "买家客户"},
        "session": {"mode": "buyer_to_target"},
        "candidates": [{
            "position": 1,
            "candidate": {
                "name": "示例标的",
                "field_groups": [{"fields": [{"label": "营收", "value": "1亿元"}]}],
            },
            "selection_snapshot": {"match_summary": "方向匹配"},
            "matching_context": {"deep_evaluation": {"reason": "建议进一步了解"}},
        }],
    }
    markdown = build_fallback_report_markdown(context, title="推荐标的报告")
    assert "| 评估维度 | 客户关注重点 | 候选现有情况 | 初步判断 |" in markdown
    assert "建议进一步了解" in markdown

    normalized = normalize_report_markdown(
        "```markdown\n<h1>不保留 HTML</h1>\n正文\n```",
        title="报告",
    )
    assert normalized.startswith("# 报告")
    assert "<h1>" not in normalized


def test_docx_renderer_preserves_heading_table_and_basic_styles() -> None:
    markdown = """# 推荐标的报告

> 用于初步判断是否继续跟进。

## 1. 示例标的

| 评估维度 | 客户关注重点 | 候选现有情况 | 初步判断 |
|---|---|---|---|
| 业务 | 技术能力 | **已有核心团队** | 可继续了解 |

- 待确认近两年客户结构
"""
    content = render_report_docx(markdown, title="推荐标的报告")
    assert content.startswith(b"PK")

    document = Document(BytesIO(content))
    assert document.core_properties.title == "推荐标的报告"
    assert document.paragraphs[0].text == "推荐标的报告"
    assert len(document.tables) == 1
    assert len(document.tables[0].columns) == 4
    assert document.tables[0].cell(1, 2).text == "已有核心团队"
    assert safe_docx_filename('推荐/报告:*?') == "推荐_报告.docx"
